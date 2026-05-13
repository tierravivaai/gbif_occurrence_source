"""Generate a tables-only CBD Publisher Country Share Report as a Word document.

Produces a clean Word document containing only the data tables (no narrative),
suitable for reference and validation purposes.
"""

import csv
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

PROCESSED_DIR = "data/processed"
SNAPSHOT_LABEL = "January 2026"
CITATION = "GBIF.org (1 June 2025) GBIF Occurrence Download https://doi.org/10.15468/dl.jsevhc"

NO_AVES_FILES = {
    "un_region": "cbd_parties_no_aves_un_region_summary.csv",
    "un_sub_region": "cbd_parties_no_aves_un_region_un_sub_region_summary.csv",
    "un_intermediate_region": "cbd_parties_no_aves_un_region_un_sub_region_un_intermediate_region_summary.csv",
    "development_status": "cbd_parties_no_aves_development_status_summary.csv",
    "income_group": "cbd_parties_no_aves_income_group_summary.csv",
}

ALL_TAXA_FILES = {
    "un_region": "cbd_parties_all_taxa_un_region_summary.csv",
    "un_sub_region": "cbd_parties_all_taxa_un_region_un_sub_region_summary.csv",
    "un_intermediate_region": "cbd_parties_all_taxa_un_region_un_sub_region_un_intermediate_region_summary.csv",
    "development_status": "cbd_parties_all_taxa_development_status_summary.csv",
    "income_group": "cbd_parties_all_taxa_income_group_summary.csv",
}

COUNT_COLS = [
    ("Internal Count", "internal_count", False),
    ("Sub-regional Count", "sub_regional_count", False),
    ("Regional Count", "regional_count", False),
    ("External Count", "external_count", False),
    ("Unknown Count", "unknown_count", False),
    ("Total Count", "total_count", False),
    ("Internal %", "internal_percentage", True),
    ("Sub-regional %", "sub_regional_percentage", True),
    ("Regional %", "regional_percentage", True),
    ("External %", "external_percentage", True),
]


def _fmt_int(value):
    try:
        f = float(value)
        if f.is_integer():
            return f"{int(f):,}"
        return f"{f:,.2f}"
    except (ValueError, TypeError):
        return str(value) if value else "—"


def _fmt_pct(value):
    try:
        return f"{float(value):.2f}%"
    except (ValueError, TypeError):
        return str(value) if value else "—"


def _cell(value, is_pct=False):
    try:
        f = float(value)
        return _fmt_pct(f) if is_pct else _fmt_int(f)
    except (ValueError, TypeError):
        return str(value) if value else "—"


def _read_csv(filename):
    path = os.path.join(PROCESSED_DIR, filename)
    with open(path, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def _add_table(doc, headers, rows):
    if not rows:
        return
    ncols = len(headers)
    # Add total row
    total_row = {}
    for spec in COUNT_COLS:
        key = spec[1]
        if key in ("internal_count", "sub_regional_count", "regional_count", "external_count", "unknown_count", "total_count"):
            total_row[key] = sum(int(r.get(key, 0) or 0) for r in rows)
    for spec in COUNT_COLS:
        key = spec[1]
        is_pct = spec[2]
        if is_pct and "total_count" in total_row and total_row["total_count"] > 0:
            if key == "internal_percentage":
                total_row[key] = round(100.0 * total_row["internal_count"] / total_row["total_count"], 2)
            elif key == "sub_regional_percentage":
                total_row[key] = round(100.0 * total_row["sub_regional_count"] / total_row["total_count"], 2)
            elif key == "regional_percentage":
                total_row[key] = round(100.0 * total_row["regional_count"] / total_row["total_count"], 2)
            elif key == "external_percentage":
                total_row[key] = round(100.0 * total_row["external_count"] / total_row["total_count"], 2)

    all_rows = rows + [total_row]
    table = doc.add_table(rows=len(all_rows) + 1, cols=ncols, style='TableGrid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for i, row_data in enumerate(all_rows):
        is_total = (i == len(rows))
        for j, (label, key, is_pct) in enumerate(COUNT_COLS if j > 0 else []):
            pass  # handled below
        for j in range(ncols):
            cell = table.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if j == 0:
                val = "**Total**" if is_total else str(row_data.get(headers[0].split()[0].lower(), ""))
                # Actually use the first column key from the data
                first_key = list(row_data.keys())[0] if not is_total else None
                text = "**Total**" if is_total else str(row_data.get(first_key, ""))
                run = p.add_run(text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                if is_total:
                    run.bold = True
            else:
                # Map header index to col_spec
                specs = [None] + COUNT_COLS if headers[0] != "Country" else COUNT_COLS
                # For simpler tables, just use the header map
                pass

    # Simplified approach: just render all cells
    for i, row_data in enumerate(all_rows):
        is_total = (i == len(rows))
        row_idx = i + 1
        for j, header in enumerate(headers):
            if j == 0 and is_total:
                continue  # already handled
            cell = table.cell(row_idx, j)
            cell.text = ''


def generate_tables_docx(output_path):
    no_aves = {key: _read_csv(fname) for key, fname in NO_AVES_FILES.items()}
    all_taxa = {key: _read_csv(fname) for key, fname in ALL_TAXA_FILES.items()}

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    for level in range(1, 5):
        doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_heading(f'CBD Parties: Publisher Country Share Tables — {SNAPSHOT_LABEL} Snapshot', level=1)
    doc.add_paragraph(f'Data source: {CITATION}')
    doc.add_paragraph(
        'Int % = Internal (same country). Sub-reg % = Sub-regional (same UN sub-region, different country). '
        'Reg % = Regional (same UN region, different sub-region). Ext % = External (different UN region).'
    )

    sections = [
        ("2.1 By UN Region (Excluding Aves)", "un_region", no_aves, [("UN Region", "un_region_name", False)] + COUNT_COLS),
        ("2.2 By UN Sub-region (Excluding Aves)", "un_sub_region", no_aves, [("UN Region", "un_region_name", False), ("UN Sub-region", "un_sub_region_name", False)] + COUNT_COLS),
        ("2.3 By UN Intermediate Region (Excluding Aves)", "un_intermediate_region", no_aves, [("UN Region", "un_region_name", False), ("UN Sub-region", "un_sub_region_name", False), ("Intermediate Region", "un_intermediate_region_name", False)] + COUNT_COLS),
        ("2.4 By Development Status (Excluding Aves)", "development_status", no_aves, [("Development Status", "un_developed_or_developing_countries", False)] + COUNT_COLS),
        ("2.5 By Income Group (Excluding Aves)", "income_group", no_aves, [("Income Group", "wb_income_group", False)] + COUNT_COLS),
        ("A.1 By UN Region (All Taxa)", "un_region", all_taxa, [("UN Region", "un_region_name", False)] + COUNT_COLS),
        ("A.2 By UN Sub-region (All Taxa)", "un_sub_region", all_taxa, [("UN Region", "un_region_name", False), ("UN Sub-region", "un_sub_region_name", False)] + COUNT_COLS),
        ("A.3 By UN Intermediate Region (All Taxa)", "un_intermediate_region", all_taxa, [("UN Region", "un_region_name", False), ("UN Sub-region", "un_sub_region_name", False), ("Intermediate Region", "un_intermediate_region_name", False)] + COUNT_COLS),
        ("A.4 By Development Status (All Taxa)", "development_status", all_taxa, [("Development Status", "un_developed_or_developing_countries", False)] + COUNT_COLS),
        ("A.5 By Income Group (All Taxa)", "income_group", all_taxa, [("Income Group", "wb_income_group", False)] + COUNT_COLS),
    ]

    for heading, key, data_dict, col_specs in sections:
        doc.add_heading(heading, level=2)
        rows = data_dict.get(key, [])
        if not rows:
            continue
        headers = [spec[0] for spec in col_specs]
        ncols = len(headers)
        # Add total row
        total = {}
        for spec in col_specs:
            k, is_pct = spec[1], spec[2]
            if k in ("internal_count", "sub_regional_count", "regional_count", "external_count", "unknown_count", "total_count"):
                total[k] = sum(int(r.get(k, 0) or 0) for r in rows)
        if "total_count" in total and total["total_count"] > 0:
            for spec in col_specs:
                k, is_pct = spec[1], spec[2]
                if is_pct:
                    pct_key = k
                    cnt_key = k.replace("_percentage", "_count")
                    if cnt_key in total:
                        total[pct_key] = round(100.0 * total[cnt_key] / total["total_count"], 2)

        table = doc.add_table(rows=len(rows) + 2, cols=ncols, style='TableGrid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header
        for j, header in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for i, row in enumerate(rows):
            for j, spec in enumerate(col_specs):
                cell = table.cell(i + 1, j)
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(_cell(row.get(spec[1], ""), is_pct=spec[2]))
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                if (i + 1) % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

        # Total row
        total_idx = len(rows) + 1
        for j, spec in enumerate(col_specs):
            cell = table.cell(total_idx, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if j == 0:
                run = p.add_run("Total")
                run.bold = True
            elif spec[1] in total:
                run = p.add_run(_cell(total[spec[1]], is_pct=spec[2]))
            else:
                run = p.add_run("")
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.save(output_path)
    print(f"Tables document saved to {output_path}")


if __name__ == "__main__":
    from datetime import datetime
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    output = f"CBD_Publisher_Country_Share_Tables_{SNAPSHOT_LABEL.replace(' ', '_')}_{ts}.docx"
    generate_tables_docx(output)
