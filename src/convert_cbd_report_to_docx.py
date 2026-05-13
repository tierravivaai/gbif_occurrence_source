"""Convert CBD_Publisher_Country_Share_Report.md to Word (.docx) using python-docx."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = "CBD_Publisher_Country_Share_Report.md"
DOCX_PATH = "CBD_Publisher_Country_Share_Report.docx"


def _add_formatted_runs(paragraph, text):
    """Parse inline markdown formatting (bold, italic, code, links) into runs."""
    # Handle links first: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    last = 0
    for m in re.finditer(link_pattern, text):
        # Add any text before the link
        before = text[last:m.start()]
        if before:
            _add_inline_runs(paragraph, before)
        # Add hyperlink
        _add_hyperlink(paragraph, m.group(1), m.group(2))
        last = m.end()
    # Add remaining text after last link
    remaining = text[last:]
    if remaining:
        _add_inline_runs(paragraph, remaining)


def _add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
    )
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_inline_runs(paragraph, text):
    """Parse bold, italic, code formatting into runs."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def _parse_table_rows(lines, start_idx):
    """Extract table rows starting at start_idx. Returns (rows_as_lists, next_line_idx)."""
    rows = []
    idx = start_idx
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(set(c) <= {'-', ':', ' '} for c in cells):
            idx += 1
            continue
        rows.append(cells)
        idx += 1
    return rows, idx


def _add_table(doc, rows):
    """Add a Word table from parsed row data. Light grey header, alternating body rows."""
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols, style='TableGrid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                continue
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            if i == 0:
                # Header row: light grey background, dark text, bold
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                # Body row: white background, black text
                run = p.add_run(cell_text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                # Alternating row shading (lighter grey for even rows)
                if i % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                    cell._tc.get_or_add_tcPr().append(shading)


def _strip_yaml_front_matter(lines):
    """Remove YAML front matter (enclosed by --- delimiters at the start)."""
    if lines and lines[0].strip() == '---':
        end = None
        for i in range(1, len(lines)):
            if lines[i].strip() == '---':
                end = i + 1
                break
        if end is not None:
            return lines[end:]
    return lines


def convert(md_path=MD_PATH, docx_path=DOCX_PATH):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
    lines = _strip_yaml_front_matter(lines)

    doc = Document()

    # Set default style to black text
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Ensure all heading styles are black
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    idx = 0
    while idx < len(lines):
        line = lines[idx]

        if not line.strip():
            idx += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            idx += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                heading = doc.add_heading(text, level=0)
            else:
                heading = doc.add_heading(text, level=level)
            idx += 1
            continue

        # Table
        if line.strip().startswith('|'):
            rows, next_idx = _parse_table_rows(lines, idx)
            _add_table(doc, rows)
            idx = next_idx
            continue

        # Unordered list
        m = re.match(r'^-\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, m.group(1))
            idx += 1
            continue

        # Ordered list
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            _add_formatted_runs(p, m.group(2))
            idx += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_runs(p, line.strip())
        idx += 1

    doc.save(docx_path)
    print(f"Word document saved to {docx_path}")


if __name__ == "__main__":
    convert()
